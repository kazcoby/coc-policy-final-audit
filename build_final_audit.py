#!/usr/bin/env python3
"""Side-by-side audit: current formatting (left) vs final accessible (right),
with a per-document accessibility/structure header. All 156 policies."""
import glob,os,re,subprocess,html,json,shutil
from collections import Counter
from docx import Document
from docx.oxml.ns import qn
HERE=os.path.dirname(os.path.abspath(__file__))
BASE=os.path.expanduser("~/Desktop/Online Ed")
CUR=os.path.join(BASE,"Policies_Upload")
FIN=os.path.join(BASE,"Final Accessible Policy Uploads")
IMG=os.path.join(HERE,"img"); os.makedirs(IMG,exist_ok=True)
SOFFICE="/opt/homebrew/bin/soffice"; RES="100"
results={ (r['sub'],r['name']):r for r in json.load(open("/tmp/sf_results.json")) }
def esc(s): return html.escape(str(s),quote=True)
def code(n):
    m=re.match(r'^([A-Z]+)\s*(\d+)',n); return (m.group(1)+m.group(2)) if m else n[:8]
IGN=set("cpc reviewed adopted approved endorsed board date next review page administrative procedure policy college of the canyons santa clarita spring fall revised".split())
ROMAN=set("i ii iii iv v vi vii viii ix x xi xii xiii xiv xv xvi xvii xviii".split())
def fulltext(path):
    d=Document(path); parts=[p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells: parts.append(cell.text)
    return " ".join(parts)
def realtoks(t): return Counter(x for x in re.findall(r"[a-z0-9]+",t.lower()) if len(x)>1 and x not in ROMAN and not x.isdigit())

docs=[]
for sub in ("3000","4000","5000"):
    for f in sorted(glob.glob(os.path.join(FIN,sub,"*.docx"))):
        docs.append((sub,f))

# migrate any bare {code}-N.png (earlier single-column render of FINAL) -> {code}_new-N.png
for p in glob.glob(os.path.join(IMG,"*.png")):
    b=os.path.basename(p)
    if re.match(r'^[A-Z]+\d+-\d+\.png$', b):
        os.rename(p, os.path.join(IMG, b.replace('-','_new-',1)))

def render(side, folder):
    if glob.glob(os.path.join(IMG,f"*_{side}-*.png")): return
    tmp=os.path.join(HERE,"_pdf_"+side); os.makedirs(tmp,exist_ok=True)
    files=[os.path.join(folder,sub,os.path.basename(f)) for sub,f in docs]
    files=[f for f in files if os.path.exists(f)]
    for k in range(0,len(files),20):
        subprocess.run([SOFFICE,"--headless","--convert-to","pdf","--outdir",tmp]+files[k:k+20],capture_output=True,timeout=600)
    for f in files:
        c=code(os.path.basename(f))
        pdf=os.path.join(tmp,os.path.splitext(os.path.basename(f))[0]+".pdf")
        if os.path.exists(pdf):
            subprocess.run(["pdftoppm","-png","-r",RES,pdf,os.path.join(IMG,c+"_"+side)],capture_output=True)
    shutil.rmtree(tmp,ignore_errors=True)
render("new", FIN)      # right column (reuses migrated images if present)
render("cur", CUR)      # left column

def pages(c,side):
    hits=glob.glob(os.path.join(IMG,f"{c}_{side}-*.png"))
    return [os.path.basename(x) for x in sorted(hits,key=lambda x:int(re.search(r'-(\d+)\.png$',x).group(1)))]
def src(name):
    # cache-busting: mtime query so browsers/CDN refetch changed images
    try: v=int(os.path.getmtime(os.path.join(IMG,name)))
    except OSError: v=0
    return f"img/{name}?v={v}"

def metrics(sub,f):
    d=Document(f)
    hs=[int(p.style.name.split()[1]) for p in d.paragraphs if p.style.name.startswith("Heading") and p.text.strip()]
    h1=hs.count(1); skip=any(hs[i]>hs[i-1]+1 for i in range(1,len(hs)))
    lang=d.styles.element.find(qn('w:docDefaults')) is not None
    alt=all((s._inline.find(qn('wp:docPr')) is not None and (s._inline.find(qn('wp:docPr')).get('descr') or '').strip()) for s in d.inline_shapes)
    acc=(h1==1 and not skip and lang and alt)
    cur=os.path.join(CUR,sub,os.path.basename(f))
    if os.path.exists(cur):
        ct=realtoks(fulltext(cur)); ft=realtoks(fulltext(f))
        diff={k:v for k in set(ct)|set(ft) if (v:=ct[k]-ft[k]) and k not in IGN}
    else: diff={}
    r=results.get((sub,os.path.basename(f)),{})
    return dict(h=(hs.count(1),hs.count(2),hs.count(3)),acc=acc,integ=(not diff),
                rlists=r.get('real_lists',0),flags=r.get('flags',[]),irr=r.get('irregular_marker_paras',0))

secs={"3000":[],"4000":[],"5000":[]}; toc={"3000":[],"4000":[],"5000":[]}
n_acc=n_flag=n_int=tot_lists=0
for sub,f in docs:
    c=code(os.path.basename(f)); m=metrics(sub,f)
    n_acc+=m['acc']; n_int+=m['integ']; tot_lists+=m['rlists']
    if m['flags'] or m['irr']: n_flag+=1
    toc[sub].append(f'<a href="#{c}" style="margin:2px 6px;color:#003366;text-decoration:none;font-size:12.5px;">{esc(c[:2]+" "+c[2:])}</a>')
    badges=(f'<span style="background:{"#1d6b3f" if m["acc"] else "#b3771a"};color:#fff;padding:2px 9px;border-radius:11px;font-size:11px;font-weight:700;">{"ACCESSIBLE" if m["acc"] else "REVIEW"}</span> '
            f'<span style="background:#eef;color:#334;padding:2px 8px;border-radius:11px;font-size:11px;">H1/H2/H3 = {m["h"][0]}/{m["h"][1]}/{m["h"][2]}</span> '
            f'<span style="background:#eef;color:#334;padding:2px 8px;border-radius:11px;font-size:11px;">{m["rlists"]} real lists</span> '
            f'<span style="background:{"#e8f3ec" if m["integ"] else "#fbeada"};color:#334;padding:2px 8px;border-radius:11px;font-size:11px;">content {"preserved" if m["integ"] else "check"}</span>'
            + (f' <span style="background:#fff3d6;color:#8a5300;padding:2px 8px;border-radius:11px;font-size:11px;">&#9873; {len(m["flags"])+(1 if m["irr"] else 0)} to review</span>' if (m['flags'] or m['irr']) else ''))
    flagnote=''
    if m['flags'] or m['irr']:
        items=m['flags'][:]
        if m['irr']: items.append(f"{m['irr']} paragraph(s) kept literal markers (paren/double-letter/single-item lists)")
        flagnote='<div style="background:#fff9ec;border-left:4px solid #d9a300;padding:8px 12px;margin:6px 14px;font-size:12px;color:#6b5300;">Needs your eyes: '+'; '.join(esc(x) for x in items)+'</div>'
    cur_p=pages(c,"cur"); new_p=pages(c,"new"); npg=max(len(cur_p),len(new_p))
    rows=[]
    for i in range(npg):
        l=(f'<img src="{src(cur_p[i])}" loading="lazy" style="width:100%;border:2px solid #b3771a;display:block;">' if i<len(cur_p) else '<div style="padding:26px;color:#9aa;text-align:center;border:2px solid #e3c9a0;">(no page)</div>')
        r=(f'<img src="{src(new_p[i])}" loading="lazy" style="width:100%;border:2px solid #1d6b3f;display:block;">' if i<len(new_p) else '<div style="padding:26px;color:#9aa;text-align:center;border:2px solid #a9cdb8;">(no page)</div>')
        rows.append(f'<tr><td width="50%" style="vertical-align:top;padding:8px;"><div style="background:#b3771a;color:#fff;font-size:11px;font-weight:700;padding:4px 9px;border-radius:5px 5px 0 0;">CURRENT &middot; p{i+1}</div>{l}</td>'
                    f'<td width="50%" style="vertical-align:top;padding:8px;"><div style="background:#1d6b3f;color:#fff;font-size:11px;font-weight:700;padding:4px 9px;border-radius:5px 5px 0 0;">FINAL ACCESSIBLE &middot; p{i+1}</div>{r}</td></tr>')
    secs[sub].append(f'<section id="{c}" style="margin:0 0 32px;border:1px solid #d7dee7;border-radius:9px;overflow:hidden;">'
        f'<div style="background:#003366;color:#fff;padding:10px 16px;position:sticky;top:0;z-index:5;"><b style="font-size:15px;">{esc(os.path.splitext(os.path.basename(f))[0])}</b>'
        f'<a href="#top" style="float:right;color:#cdd9e8;font-size:12px;text-decoration:none;">&#8679; top</a><div style="margin-top:6px;">{badges}</div></div>'
        f'{flagnote}<table width="100%" style="border-collapse:collapse;">{"".join(rows)}</table></section>')

N=len(docs)
banner=(f'<div style="background:#003366;color:#fff;border-radius:9px;padding:20px 24px;margin-bottom:14px;">'
 f'<div style="font-size:23px;font-weight:700;">Final Accessible Policy Uploads &mdash; Audit</div>'
 f'<div style="font-size:14px;color:#cdd9e8;margin-top:4px;">{N} policies &middot; current formatting (left) vs final accessible (right) &middot; College of the Canyons</div>'
 f'<div style="margin-top:12px;font-size:14px;">'
 f'<span style="background:#1d6b3f;padding:4px 12px;border-radius:12px;margin-right:6px;">{n_acc}/{N} pass accessibility</span>'
 f'<span style="background:#2c5a8a;padding:4px 12px;border-radius:12px;margin-right:6px;">{tot_lists} real lists</span>'
 f'<span style="background:#2c5a8a;padding:4px 12px;border-radius:12px;margin-right:6px;">{n_int}/{N} content preserved</span>'
 f'<span style="background:#8a6a00;padding:4px 12px;border-radius:12px;">{n_flag} need review</span></div></div>')
tocblock="".join(f'<div style="background:#fff;border-radius:9px;padding:12px 18px;margin-bottom:10px;"><b style="color:#003366;">{s} series ({len(toc[s])})</b><div style="margin-top:6px;line-height:1.9;">{"".join(toc[s])}</div></div>' for s in ("3000","4000","5000"))
body="".join(secs["3000"]+secs["4000"]+secs["5000"])
page=f'''<!DOCTYPE html><html lang="en"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>Final Accessible Policy Uploads - Audit</title></head>
<body style="margin:0;background:#eef1f5;font-family:'Segoe UI',Arial,sans-serif;color:#1a2733;" id="top"><div style="max-width:1200px;margin:0 auto;padding:22px 16px;">{banner}{tocblock}{body}</div></body></html>'''
open(os.path.join(HERE,"index.html"),"w").write(page)
print(f"built side-by-side audit: {N} docs, accessible {n_acc}/{N}, content preserved {n_int}/{N}, flagged {n_flag}")
